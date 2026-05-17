"""Initial EDA for SAEM 2026 Hackathon dataset.

Reads:
  1. Hackathon_Codebook_Release_1_SHARE.docx — variable definitions
  2. Hackathon_Data_Release_1_SHARE.xlsx — three sheets

Reports shapes, dtypes, missingness, label balance, and triage-note
length distributions. Aggregate-only output; no row data printed.
"""
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

DATA_DIR = Path(__file__).resolve().parents[2] / "SAEM_2026_hackathon" / "data"
CODEBOOK = DATA_DIR / "Hackathon_Codebook_Release_1_SHARE.docx"
XLSX = DATA_DIR / "Hackathon_Data_Release_1_SHARE.xlsx"


def section(title: str) -> None:
    print(f"\n{'=' * 78}\n{title}\n{'=' * 78}")


def dump_codebook(path: Path) -> None:
    section(f"CODEBOOK — {path.name}")
    doc = Document(str(path))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    print(f"Paragraphs (non-empty): {len(paragraphs)}")
    print(f"Tables: {len(doc.tables)}")

    print("\n--- All paragraphs ---")
    for p in paragraphs:
        print(p)

    for i, table in enumerate(doc.tables, 1):
        print(f"\n--- Table {i} ({len(table.rows)} rows × "
              f"{len(table.columns)} cols) ---")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            print(" | ".join(cells))


def dump_sheet(name: str, df: pd.DataFrame) -> None:
    section(f"SHEET — {name}  shape={df.shape}")

    print("\n--- Columns + dtypes ---")
    for col in df.columns:
        print(f"  {col:40s}  {str(df[col].dtype):15s}")

    print("\n--- Missingness (count, pct) ---")
    miss = df.isna().sum()
    n = len(df)
    for col in df.columns:
        c = int(miss[col])
        if c > 0:
            print(f"  {col:40s}  {c:>6d}  {c / n * 100:5.1f}%")
    if (miss == 0).all():
        print("  (no missing values in any column)")

    print("\n--- Numeric columns: describe ---")
    num = df.select_dtypes(include="number")
    if num.shape[1] > 0:
        with pd.option_context("display.width", 200, "display.max_columns", 50):
            print(num.describe().T[["count", "mean", "std", "min",
                                    "25%", "50%", "75%", "max"]])
    else:
        print("  (no numeric columns)")

    print("\n--- Categorical / object columns: nunique + top values ---")
    obj = df.select_dtypes(include="object")
    for col in obj.columns:
        nunique = df[col].nunique(dropna=True)
        top = df[col].value_counts(dropna=False).head(8)
        print(f"\n  Column: {col}  (nunique={nunique})")
        for val, cnt in top.items():
            display_val = (
                "<NA>" if pd.isna(val)
                else (val[:60] + "...") if isinstance(val, str) and len(val) > 60
                else val
            )
            print(f"    {str(display_val):65s}  {cnt}")

        # Length distribution for likely text columns
        sample_vals = df[col].dropna().astype(str)
        if sample_vals.empty:
            continue
        avg_len = sample_vals.str.len().mean()
        if avg_len > 50:  # likely a note / narrative column
            lens = sample_vals.str.len()
            words = sample_vals.str.split().str.len()
            print(f"    [text-like] char_len:  "
                  f"min={lens.min()}, p25={lens.quantile(.25):.0f}, "
                  f"med={lens.median():.0f}, p75={lens.quantile(.75):.0f}, "
                  f"max={lens.max()}, mean={lens.mean():.0f}")
            print(f"                word_count: "
                  f"min={words.min()}, p25={words.quantile(.25):.0f}, "
                  f"med={words.median():.0f}, p75={words.quantile(.75):.0f}, "
                  f"max={words.max()}, mean={words.mean():.0f}")


def main() -> None:
    section("FILE OVERVIEW")
    print(f"Codebook: {CODEBOOK}  exists={CODEBOOK.exists()}  "
          f"size={CODEBOOK.stat().st_size if CODEBOOK.exists() else 'N/A'} bytes")
    print(f"XLSX:     {XLSX}  exists={XLSX.exists()}  "
          f"size={XLSX.stat().st_size if XLSX.exists() else 'N/A'} bytes")

    if CODEBOOK.exists():
        dump_codebook(CODEBOOK)

    if XLSX.exists():
        section(f"XLSX — {XLSX.name}")
        xl = pd.ExcelFile(XLSX, engine="openpyxl")
        print(f"Sheets: {xl.sheet_names}")
        for name in xl.sheet_names:
            df = xl.parse(name)
            dump_sheet(name, df)


if __name__ == "__main__":
    main()
